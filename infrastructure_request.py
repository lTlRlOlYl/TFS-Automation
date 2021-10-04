from docx import Document
import os

class AuditInfrastructureRequest():
    inf_req_template = r'word\inf_req_template.docx'
    output_path=None
    variables=None

    def __init__(
        self,
        agency_acronym,
        audit_number,
        audit_name,
        analyst_tag,
        analyst_name,
        ):
        self.agency_acronym = agency_acronym
        self.audit_number = audit_number
        self.audit_number_nohyph = audit_number.replace("-","")
        self.audit_name = audit_name
        self.analyst_tag = analyst_tag
        self.analyst_name = analyst_name
        self.set_variables()
        self.set_output_path()


    #called from __init__ to set the variables dict the create_inf_req_doc iterates over 
    def set_variables(self):
        self.variables = {
            "${DB_NAME}": "Audit_{agency_acronym}_{audit_number_nohyph}".format(
                agency_acronym=self.agency_acronym, 
                audit_number_nohyph=self.audit_number_nohyph
                ),
            "${AUDIT_NAME}": "{audit_name}".format(audit_name=self.audit_name),
            "${ANALYST_NAME}": "{analyst_name}".format(analyst_name=self.analyst_name)
        }


    #called from __init__ to specify where the inf req doc is be saved
    def set_output_path(self):
        self.output_path = r"word\{db_name} Infrastructure Request.docx".format(
            db_name=self.variables['${DB_NAME}']
            )


    #helper function used in create_inf_req()
    def replace_text_in_paragraph(self, paragraph, key, value):
        if key in paragraph.text:
            inline = paragraph.runs
            for item in inline:
                if key in item.text:
                    item.text = item.text.replace(key, value)


    #opens a word doc based on a template in project dir, inserts instance data, and saves
    def create_inf_req_doc(self):
        template_document = Document(self.inf_req_template)

        for variable_key, variable_value in self.variables.items():
            for paragraph in template_document.paragraphs:
                self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(self.output_path)
        return self.output_path
    

    #deletes the inf req doc
    def delete_inf_req_doc(self):
        if os.path.exists(self.output_path):
            os.remove(self.output_path)